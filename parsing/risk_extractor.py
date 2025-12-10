# risk_extractor.py
import json
import re
import os
from typing import List, Dict, Optional
from docx import Document

class StrictRuleBasedExtractor:
    def __init__(self):
        self.current_document_type = "unknown"
    
    def extract_risks(self, filepath: str, doc_type: str) -> List[Dict]:
        print(f"извлечение рисков из: {os.path.basename(filepath)}")
        
        if not os.path.exists(filepath):
            print(f"файл не найден: {filepath}")
            return []
        
        self.current_document_type = doc_type
        
        if "services" in doc_type:
            return self.extract_services_risks(filepath, doc_type)
        else:
            return self.extract_standard_risks(filepath, doc_type)
    
    def extract_services_risks(self, filepath: str, doc_type: str) -> List[Dict]:
        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"ошибка чтения docx: {e}")
            return []
        
        all_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_paragraphs.append(text)
        
        risks = []
        current_header = None
        current_content = []
        
        for i in range(len(all_paragraphs)):
            text = all_paragraphs[i]
            
            if self.is_service_risk_header(text):
                if current_header and current_content:
                    content = ' '.join(current_content).strip()
                    if content and len(content) > 30:
                        risk = self.create_risk(current_header, content, len(risks), doc_type, filepath)
                        risks.append(risk)
                
                current_header = text
                current_content = []
            
            elif current_header is not None:
                if i + 1 < len(all_paragraphs) and self.is_service_risk_header(all_paragraphs[i + 1]):
                    current_content.append(text)
                    content = ' '.join(current_content).strip()
                    if content and len(content) > 30:
                        risk = self.create_risk(current_header, content, len(risks), doc_type, filepath)
                        risks.append(risk)
                    current_header = None
                    current_content = []
                else:
                    current_content.append(text)
        
        if current_header and current_content:
            content = ' '.join(current_content).strip()
            if content and len(content) > 30:
                risk = self.create_risk(current_header, content, len(risks), doc_type, filepath)
                risks.append(risk)
        
        print(f"  найдено рисков: {len(risks)}")
        return risks
    
    def is_service_risk_header(self, text: str) -> bool:
        if not text:
            return False
        
        text_lower = text.lower()
        
        if self.is_general_header(text):
            return False
        
        if re.match(r'^\d+\.\s+[а-я]', text_lower):
            if len(text) < 100:
                return True
        
        return False
    
    def extract_standard_risks(self, filepath: str, doc_type: str) -> List[Dict]:
        try:
            doc = Document(filepath)
        except Exception as e:
            print(f"ошибка чтения docx: {e}")
            return []
        
        all_paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_paragraphs.append(text)
        
        risks = []
        
        for i in range(len(all_paragraphs)):
            text = all_paragraphs[i]
            if self.is_standard_risk_header(text):
                print(f"  найден риск: {text[:60]}...")
                
                header = text
                content_parts = []
                
                j = i + 1
                while j < len(all_paragraphs):
                    next_text = all_paragraphs[j]
                    
                    if self.is_standard_risk_header(next_text):
                        break
                    
                    if self.is_general_header(next_text):
                        break
                    
                    content_parts.append(next_text)
                    j += 1
                
                content = ' '.join(content_parts)
                
                if content.strip() and len(content) > 30:
                    risk = self.create_risk(header, content, len(risks), doc_type, filepath)
                    risks.append(risk)
        
        print(f"  найдено рисков: {len(risks)}")
        return risks
    
    def is_standard_risk_header(self, text: str) -> bool:
        patterns = [
            r'^\d+\.\d+\.\s+',
            r'^\d+\.\s+риск\s+',
            r'^риск\s+\w+\s+при',
        ]
        
        for pattern in patterns:
            if re.match(pattern, text, re.IGNORECASE):
                if not self.is_general_header(text):
                    return True
        
        return False
    
    def is_general_header(self, text: str) -> bool:
        text_lower = text.lower()
        
        general_patterns = [
            'оглавление:',
            'риски поставщика при заключении',
            'риски покупателя при заключении',
            'риски подрядчика при заключении',
            'рассмотрим основные',
            'при заключении договора',
            'если предмет договора',
            'договор возмездного оказания услуг',
            'материал позволяет',
            'проверить другие варианты',
            'проверить налоговые риски',
            'узнать больше деталей',
            'скачать готовую форму',
            'готовое решение:',
            'консультантплюс',
            'документ предоставлен',
            'дата сохранения:',
            'актуально на',
            'см. аналогичный',
        ]
        
        return any(pattern in text_lower for pattern in general_patterns)
    
    def create_risk(self, header: str, text: str, idx: int, doc_type: str, filepath: str) -> Dict:
        articles = self.extract_articles_with_context(text)
        
        consequences = self.extract_consequences(text)
        
        severity = self.estimate_severity(text)
        
        category = self.determine_category(header, text)
        
        risk_title = self.clean_risk_title(header)
        
        description = self.create_description(text)
        
        recommendation = self.extract_recommendation(text)
        
        pattern = self.extract_pattern(header)
        
        risk_number = self.extract_risk_number(header)
        
        return {
            "id": f"risk_{doc_type}_{idx}",
            "type": "risk",
            "header": header,
            "body": text[:500] + "..." if len(text) > 500 else text,
            "summary": text[:200] + "..." if len(text) > 200 else text,
            "metadata": {
                "risk_title": risk_title,
                "description": description,
                "pattern": pattern,
                "recommendation": recommendation,
                "severity": severity,
                "consequences": consequences,
                "relevant_articles": articles,
                "risk_category": category,
                "document_type": doc_type,
                "source_file": os.path.basename(filepath),
                "risk_number": risk_number,
                "source": "консультантплюс, 2025",
                "extraction_method": "strict rule-based"
            }
        }
    
    def extract_articles_with_context(self, text: str) -> List[str]:
        articles = []
        
        patterns = [
            (r'ст\.\s*(\d+(?:[а-я])?(?:\s*[-–]\s*\d+[а-я]?)?)', 'ст. '),
            (r'стать[яи]\s*(\d+(?:[а-я])?(?:\s*[-–]\s*\d+[а-я]?)?)', 'статья '),
            (r'п\.\s*\d+\s*ст\.\s*(\d+)', 'ст. '),
            (r'пункт[о\w]*\s*\d+\s*стать[ия]\s*(\d+)', 'ст. '),
        ]
        
        for pattern, prefix in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                article_num = match.group(1)
                article_ref = f"{prefix}{article_num}"
                
                start = max(0, match.start() - 100)
                end = min(len(text), match.end() + 100)
                context = text[start:end].strip()
                
                if 'гк' in context.lower() or 'гражданск' in context.lower():
                    article_ref += " гк рф"
                
                if article_ref not in articles:
                    articles.append(article_ref)
        
        unique_articles = []
        for article in articles:
            if article not in unique_articles:
                unique_articles.append(article)
        
        return unique_articles[:10]
    
    def clean_risk_title(self, header: str) -> str:
        patterns = [
            r'^\d+\.\d+\.\s+',
            r'^\d+\.\s+',
            r'^риск\s+',
            r'^риски\s+'
        ]
        
        title = header
        for pattern in patterns:
            title = re.sub(pattern, '', title, flags=re.IGNORECASE)
        
        return title.strip()
    
    def extract_risk_number(self, header: str) -> str:
        match = re.search(r'^(\d+\.\d+)\.', header)
        if match:
            return match.group(1)
        
        match = re.search(r'^(\d+)\.\s+', header)
        if match:
            return match.group(1)
        
        match = re.search(r'^(\d+)\.\s+\w', header)
        if match:
            return match.group(1)
        
        return None
    
    def create_description(self, text: str) -> str:
        sentences = re.split(r'[.!?]\s+', text)
        
        for sentence in sentences:
            sentence_lower = sentence.lower()
            if any(keyword in sentence_lower for keyword in [
                'риск заключается', 'опасность состоит', 'проблема в том',
                'может привести', 'влечет', 'приводит к', 'заключается в'
            ]):
                return sentence.strip()[:150]
        
        for sentence in sentences:
            if len(sentence) > 30 and 'риск' not in sentence.lower():
                return sentence.strip()[:150]
        
        if sentences and len(sentences[0]) > 20:
            return sentences[0].strip()[:150]
        
        return "риск связан с договорными условиями"
    
    def extract_consequences(self, text: str) -> List[str]:
        consequences = []
        text_lower = text.lower()
        
        consequence_map = {
            'договор может быть признан незаключенным': 'признание договора незаключенным',
            'признан незаключенным': 'признание договора незаключенным',
            'не сможете требовать': 'потеря права требовать исполнения',
            'не вправе требовать': 'потеря права требовать исполнения',
            'обязаны будете оплатить': 'обязанность оплатить работы/товар',
            'должны будете оплатить': 'обязанность оплатить работы/товар',
            'возместить убытки': 'обязанность возместить убытков',
            'нести ответственность': 'привлечение к ответственности',
            'не сможете отказаться': 'потеря права отказаться от договора',
            'лишаетесь права': 'потеря определенного права',
            'признан недействительным': 'признание договора недействительным',
            'расторжение договора': 'расторжение договора',
            'уплатить неустойку': 'обязанность уплатить неустойку',
            'штраф': 'уплата штрафа',
            'пени': 'начисление пеней',
            'административные штрафы': 'административная ответственность',
            'признание договора трудовым': 'признание договора трудовым',
            'соблюдать трудовое законодательство': 'обязанность соблюдать трудовое законодательство',
        }
        
        for phrase, consequence in consequence_map.items():
            if phrase in text_lower:
                consequences.append(consequence)
        
        unique_consequences = []
        for c in consequences:
            if c not in unique_consequences:
                unique_consequences.append(c)
        
        if not unique_consequences:
            sentences = re.split(r'[.!?]\s+', text)
            for sentence in sentences:
                if any(word in sentence.lower() for word in ['приведет', 'повлечет', 'вызовет', 'стать причиной']):
                    unique_consequences.append(sentence.strip()[:100])
                    break
            
            if not unique_consequences:
                unique_consequences.append('негативные юридические последствия')
        
        return unique_consequences[:5]
    
    def estimate_severity(self, text: str) -> int:
        text_lower = text.lower()
        
        high_severity_phrases = [
            'договор может быть признан незаключенным',
            'расторжение договора',
            'признан незаключенным',
            'признан недействительным',
            'существенное нарушение',
            'административная ответственность',
            'уголовная ответственность',
            'ликвидация юридического лица',
            'признание договора трудовым',
            'штраф до 100 000 руб',
        ]
        
        medium_severity_phrases = [
            'неустойка',
            'штраф',
            'пени',
            'возмещение убытков',
            'отказ от оплаты',
            'отказ от приемки',
            'взыскание',
            'проценты',
            'ндфл',
        ]
        
        low_severity_phrases = [
            'рекомендуется',
            'следует',
            'необходимо',
            'желательно',
            'может потребовать',
            'вправе требовать'
        ]
        
        severity = 5
        
        for phrase in high_severity_phrases:
            if phrase in text_lower:
                severity = 9
                break
        
        if severity == 5:
            for phrase in medium_severity_phrases:
                if phrase in text_lower:
                    severity = 7
                    break
        
        if severity == 5:
            for phrase in low_severity_phrases:
                if phrase in text_lower:
                    severity = 3
                    break
        
        return severity
    
    def determine_category(self, header: str, content: str) -> str:
        header_lower = header.lower()
        content_lower = content.lower()
        
        categories = [
            ('предмет договора', ['предмет', 'наименование', 'количество', 'характеристики', 'параметры', 'объем', 'услуг']),
            ('качество товара/работ', ['качество', 'недостатки', 'дефекты', 'гаранти', 'брак', 'несоответствие']),
            ('сроки', ['срок', 'время', 'период', 'дата', 'просрочк', 'начальн', 'конечн', 'времени']),
            ('цена и оплата', ['цена', 'оплата', 'стоимость', 'платеж', 'ндс', 'предоплат', 'аванс', 'расчет', 'вознаграждение']),
            ('доставка и поставка', ['доставка', 'поставка', 'транспорт', 'перевозка', 'отгрузка']),
            ('приемка и передача', ['приемка', 'акт', 'проверка', 'осмотр', 'передача', 'сдача']),
            ('ответственность сторон', ['ответственность', 'неустойка', 'штраф', 'пени', 'убытки', 'возмещение']),
            ('расторжение договора', ['расторжение', 'отказ', 'прекращен', 'аннулирован']),
            ('имущественные права', ['собственность', 'право', 'залог', 'обременен', 'имущество']),
            ('субподряд и третьи лица', ['субподряд', 'субподрядчик', 'третье лицо', 'субисполнитель']),
            ('гарантия', ['гарантия', 'гарантийный', 'обслуживание', 'ремонт', 'замена']),
            ('налоги', ['ндфл', 'налог', 'налоги', 'налоговый']),
            ('трудовые отношения', ['трудовой', 'трудовые', 'отпуск', 'больничн', 'работник']),
        ]
        
        for category, keywords in categories:
            if any(keyword in header_lower for keyword in keywords):
                return category
        
        for category, keywords in categories:
            if any(keyword in content_lower for keyword in keywords):
                return category
        
        if self.current_document_type == "supplier":
            return "риски поставщика"
        elif self.current_document_type == "customer":
            return "риски покупателя"
        elif self.current_document_type == "contractor":
            return "риски подрядчика"
        elif "services" in self.current_document_type:
            return "риски услуг"
        
        return 'прочие риски'
    
    def extract_pattern(self, header: str) -> str:
        pattern = re.sub(r'^\d+\.\d+\.\s+', '', header)
        pattern = re.sub(r'^\d+\.\s+', '', pattern)
        pattern = re.sub(r'^риск\s+\w+\s+при\s+', '', pattern, flags=re.IGNORECASE)
        pattern = re.sub(r'^риски\s+\w+\s+при\s+', '', pattern, flags=re.IGNORECASE)
        
        return pattern.strip()
    
    def extract_recommendation(self, text: str) -> str:
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['рекомендуется', 'следует', 'необходимо', 'нужно', 'желательно']):
            sentences = re.split(r'[.!?]\s+', text)
            for sentence in sentences:
                if any(word in sentence.lower() for word in ['рекомендуется', 'следует', 'необходимо', 'нужно']):
                    rec = sentence.strip()
                    rec = re.sub(r'^(рекомендуется|следует|необходимо|нужно)[\s,:]+', '', rec, flags=re.IGNORECASE)
                    return rec[:120]
        
        if 'лучше' in text_lower or 'предпочтительнее' in text_lower:
            sentences = re.split(r'[.!?]\s+', text)
            for sentence in sentences:
                if 'лучше' in sentence.lower():
                    return sentence.strip()[:120]
        
        if 'не согласован' in text_lower:
            return 'четко согласовать условие в письменной форме'
        elif 'не указан' in text_lower or 'не определен' in text_lower:
            return 'указать конкретные параметры и характеристики в договоре'
        elif 'отсутствует' in text_lower:
            return 'включить недостающее условие в договоре'
        elif 'может быть признан незаключенным' in text_lower:
            return 'согласовать все существенные условия договора'
        elif 'признание договора трудовым' in text_lower:
            return 'избегать признаков трудовых отношений в договоре'
        
        return 'провести правовой анализ договорного условия'

def extract_risks_from_all_files():
    extractor = StrictRuleBasedExtractor()
    
    all_risks = []
    
    risk_files = [
        ("готовое решение_ риски поставщика при заключении договора по.docx", "supplier"),
        ("готовое решение_ риски покупателя при заключении договора по.docx", "customer"),
        ("готовое решение_ риски подрядчика при заключении договора по.docx", "contractor"),
        ("готовое решение_ договор возмездного оказания услуг физлицом.docx", "services_individual"),
        ("готовое решение_ договор возмездного оказания услуг между юр.docx", "services_legal")
    ]
    
    total_risks = 0
    
    for filepath, doc_type in risk_files:
        if not os.path.exists(filepath):
            print(f"файл не найден: {filepath}")
            continue
        
        print(f"\nобработка файла: {os.path.basename(filepath)}")
        
        risks = extractor.extract_risks(filepath, doc_type)
        
        for i, risk in enumerate(risks):
            risk['id'] = f"risk_{doc_type}_{i}"
            risk['metadata']['document_type'] = doc_type
            risk['metadata']['source_file'] = os.path.basename(filepath)
        
        all_risks.extend(risks)
        total_risks += len(risks)
        print(f"извлечено рисков: {len(risks)}")
    
    print(f"\nвсего извлечено рисков: {total_risks}")
    
    os.makedirs('data', exist_ok=True)
    output_file = 'data/all_risks_extracted.json'
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_risks, f, ensure_ascii=False, indent=2)
    
    print(f"\nвсе риски сохранены в: {output_file}")
    
    return all_risks
