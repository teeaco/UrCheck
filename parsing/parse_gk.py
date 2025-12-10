# parse_gk.py
import re
from typing import List, Dict
import os
import json

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("@@@@@Библиотека python-docx не установлена. Установите: pip install python-docx")

def parse_gk_file(filepath: str):
    
    if not DOCX_AVAILABLE:
        print("      Библиотека python-docx не доступна")
        return []
    
    try:
        doc = Document(filepath)
    except FileNotFoundError:
        print(f"Файл {filepath} не найден")
        return []
    except Exception as e:
        print(f"Ошибка открытия файла: {e}")
        return []
    
    articles = []
    current_article = None
    current_text = []
    
    print(f"Загружен документ с {len(doc.paragraphs)} параграфами")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Проверяем, начинается ли параграф с "Статья"
        if text.lower().startswith('статья') or re.match(r'^Статья\s+\d+', text):
            # сохраняем предыдущую статью
            if current_article is not None and current_text:
                article_data = process_article(current_article, ' '.join(current_text))
                if article_data:
                    articles.append(article_data)
            
            # новую статью
            current_article = text
            current_text = []
        elif current_article is not None:
            # текст к текущей статье
            current_text.append(text)
    
    # последнюю статью
    if current_article is not None and current_text:
        article_data = process_article(current_article, ' '.join(current_text))
        if article_data:
            articles.append(article_data)
    
    print(f" Найдено {len(articles)} статей")
    return articles

def process_article(header: str, body: str) -> Dict:
    # номер статьи
    article_match = re.search(r'Статья\s+(\d+)', header, re.IGNORECASE)
    if not article_match:
        # на всякий другой формат
        article_match = re.search(r'^(\d+)\.', header)
        if not article_match:
            return None
    
    article_num = article_match.group(1)
    
    # очищаем текст
    clean_body = re.sub(r'\s+', ' ', body)
    clean_body = re.sub(r'\([^)]*редакции[^)]*\)', '', clean_body, flags=re.IGNORECASE)
    clean_body = re.sub(r'\([^)]*дополнен[^)]*\)', '', clean_body, flags=re.IGNORECASE)
    
    if len(clean_body) < 10:
        return None
    
    # ключевые слова
    keywords = []
    text_lower = clean_body.lower()
    
    if "подряд" in text_lower:
        keywords.append("подряд")
    if any(word in text_lower for word in ["качество", "недостатки", "дефекты", "гаранти"]):
        keywords.append("качество")
    if any(word in text_lower for word in ["срок", "время", "период", "дата"]):
        keywords.append("сроки")
    if any(word in text_lower for word in ["цена", "оплата", "стоимость", "платеж"]):
        keywords.append("оплата")
    if any(word in text_lower for word in ["ответственность", "штраф", "неустойка", "пеня"]):
        keywords.append("ответственность")
    if any(word in text_lower for word in ["риск", "гибель", "повреждение"]):
        keywords.append("риски")
    if any(word in text_lower for word in ["приемка", "акт", "проверка"]):
        keywords.append("приемка")
    
    summary = clean_body[:200].strip()
    if len(clean_body) > 200:
        summary += "..."
    
    return {
        "id": f"gk_{article_num}",
        "type": "norm",
        "header": header[:100],
        "text": clean_body,
        "summary": summary,
        "metadata": {
            "source": "Гражданский кодекс РФ",
            "article": article_num,
            "keywords": keywords,
            "law_type": "civil_code",
            "last_updated": "2025"
        }
    }

def parse_all_gk_files(files: List[str] = None):
    if files is None: #ищем файлы
        files = []
        for file in os.listdir('.'):
            if file.lower().startswith('гк') and file.lower().endswith('.docx'):
                files.append(file)
    
    if not files:
        print(" Не найдено файлов ГК РФ в текущей директории!")
        print("   Убедитесь, что файлы называются как: ГК_РФ_часть1.docx, ГК_РФ_часть2.docx и т.д.")
        return []
    
    all_articles = []
    
    print(f"\n ПАРСИНГ ВСЕХ ФАЙЛОВ ГК РФ ({len(files)} файлов)")
    print("=" * 60)
    
    for filepath in files:
        print(f"\n Чтение файла: {filepath}")
        
        articles = parse_gk_file(filepath)
        all_articles.extend(articles)
        
        print(f"   Добавлено статей: {len(articles)}")
    
    print(f"\n ВСЕГО собрано статей: {len(all_articles)}")
    
    # чек на дубликаты
    unique_articles = {}
    for article in all_articles:
        article_num = article['metadata']['article']
        unique_articles[article_num] = article
    
    print(f"   Уникальных статей: {len(unique_articles)}")
    
    sorted_articles = sorted(
        unique_articles.values(), 
        key=lambda x: int(re.search(r'\d+', x['metadata']['article']).group())
    )
    
    # сохранение
    os.makedirs('data', exist_ok=True)
    output_file = 'data/gk_all_articles.json'
    
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(sorted_articles, f, ensure_ascii=False, indent=2)
    
    print(f"\nВсе статьи сохранены в: {output_file}")
    
    keyword_stats = {}
    for article in sorted_articles:
        for keyword in article['metadata']['keywords']:
            keyword_stats[keyword] = keyword_stats.get(keyword, 0) + 1
    
    print("\nСТАТИСТИКА ПО КЛЮЧЕВЫМ СЛОВАМ:")
    for keyword, count in sorted(keyword_stats.items(), key=lambda x: x[1], reverse=True):
        print(f"   {keyword}: {count}")
    
    return sorted_articles

def parse_gk(filepath: str):
    """Умный парсер, который определяет формат файла"""
    if filepath.lower().endswith('.docx'):
        return parse_gk_file(filepath)
    else:
        print(f" Неизвестный формат файла: {filepath}")
        print("   Попробую как DOCX...")
        try:
            return parse_gk_file(filepath)
        except:
            print("Ошибка парсинга файла")
            return []

#  для быстрого тестирования
def test_parse_all_gk():
    """Тестирует парсинг всех файлов ГК"""
    print("ТЕСТИРОВАНИЕ ПАРСИНГА ВСЕХ ФАЙЛОВ ГК")
    print("=" * 60)
    
    articles = parse_all_gk_files()
    
    if articles:
        print(f"\n📋 ПРИМЕРЫ НАЙДЕННЫХ СТАТЕЙ:")
        for i, article in enumerate(articles[:5]):
            print(f"\n{i+1}. Статья {article['metadata']['article']}")
            print(f"   Заголовок: {article['header'][:80]}...")
            print(f"   Кратко: {article['summary'][:100]}...")
            print(f"   Ключевые слова: {', '.join(article['metadata']['keywords'])}")
    
    return articles

