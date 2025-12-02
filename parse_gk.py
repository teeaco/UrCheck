# parse_gk.py
import re
from typing import List, Dict
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("@@@@@Библиотека python-docx не установлена. Установите: pip install python-docx")

def parse_gk_file(filepath: str):
    
    if not DOCX_AVAILABLE:
        print("      Библиотека python-docx не доступна")
        return []
    
    try:
        doc = Document(filepath)
    except FileNotFoundError:
        print(f"Файл {filepath} не найден")
        return []
    except Exception as e:
        print(f"Ошибка открытия файла: {e}")
        return []
    
    articles = []
    current_article = None
    current_text = []
    
    print(f"Загружен документ с {len(doc.paragraphs)} параграфами")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Проверяем, начинается ли параграф с "Статья"
        if text.lower().startswith('статья') or re.match(r'^Статья\s+\d+', text):
            # сохраняем предыдущую статью
            if current_article is not None and current_text:
                article_data = process_article(current_article, ' '.join(current_text))
                if article_data:
                    articles.append(article_data)
            
            # новую статью
            current_article = text
            current_text = []
        elif current_article is not None:
            # текст к текущей статье
            current_text.append(text)
        # else нахуй
    
    #  последнюю статью
    if current_article is not None and current_text:
        article_data = process_article(current_article, ' '.join(current_text))
        if article_data:
            articles.append(article_data)
    
    print(f" Найдено {len(articles)} статей")
    return articles

def process_article(header: str, body: str) -> Dict:
    #  номер статьи
    article_match = re.search(r'Статья\s+(\d+)', header, re.IGNORECASE)
    if not article_match:
        # на всякий другой формат
        article_match = re.search(r'^(\d+)\.', header)
        if not article_match:
            return None
    
    article_num = article_match.group(1)
    
    # очищаем текст
    clean_body = re.sub(r'\s+', ' ', body)
    clean_body = re.sub(r'\([^)]*редакции[^)]*\)', '', clean_body, flags=re.IGNORECASE)
    clean_body = re.sub(r'\([^)]*дополнен[^)]*\)', '', clean_body, flags=re.IGNORECASE)
    
    if len(clean_body) < 10:
        return None
    # ключевые слова
    keywords = []
    text_lower = clean_body.lower()
    
    if "подряд" in text_lower:
        keywords.append("подряд")
    if any(word in text_lower for word in ["качество", "недостатки", "дефекты", "гаранти"]):
        keywords.append("качество")
    if any(word in text_lower for word in ["срок", "время", "период", "дата"]):
        keywords.append("сроки")
    if any(word in text_lower for word in ["цена", "оплата", "стоимость", "платеж"]):
        keywords.append("оплата")
    if any(word in text_lower for word in ["ответственность", "штраф", "неустойка", "пеня"]):
        keywords.append("ответственность")
    if any(word in text_lower for word in ["риск", "гибель", "повреждение"]):
        keywords.append("риски")
    if any(word in text_lower for word in ["приемка", "акт", "проверка"]):
        keywords.append("приемка")
    
    summary = clean_body[:200].strip()
    if len(clean_body) > 200:
        summary += "..."
    
    return {
        "id": f"gk_{article_num}",
        "type": "norm",
        "header": header[:100],  #заголовок для отладки
        "text": clean_body,
        "summary": summary,
        "metadata": {
            "source": "Гражданский кодекс РФ",
            "article": article_num,
            "keywords": keywords,
            "law_type": "civil_code",
            "last_updated": "2025"  # можно извлечь из документа (нахуй) или указать текущую дату
        }
    }


def parse_gk(filepath: str):
    """Умный парсер, который определяет формат файла"""
    if filepath.lower().endswith('.docx'):
        return parse_gk_file(filepath)
    else:
        print(f" Неизвестный формат файла: {filepath}")
        print("   Попробую как DOCX...")
        try:
            return parse_gk_file(filepath)
        except:
            print("хуйня какая-то...")